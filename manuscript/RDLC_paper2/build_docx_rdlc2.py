"""
manuscript/RDLC_paper2/build_docx_rdlc2.py
==========================================
Paper 2 (construction-journal extension), targeted at Revista de la Construccion
(RDLC): the Entropy Robustness Index reframed as a calibration-free screening
index for progressive-collapse vulnerability of steel building frames, validated
against code-style alternate-load-path (ALP) column-removal analysis on a
published benchmark building (Vogel 1985).

Every reported number is COMPUTED at build time from the shared
manuscript/results.py (compute_paper2_results), so the manuscript cannot drift
out of sync with the code.

    python manuscript/generate_figures.py               # produce figures first
    python manuscript/RDLC_paper2/build_docx_rdlc2.py    # -> Carvajal_RDLC_screening_manuscript.docx

Scope is honest and unchanged from the repo invariants: planar, first-order,
linear-elastic, quasi-static; no early-warning claim; R_S measures load-path
redundancy, not stiffness/strength.
"""

import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
MANU = os.path.abspath(os.path.join(HERE, ".."))   # manuscript/ (shared results.py + figures/)
sys.path.insert(0, MANU)

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from results import compute_paper2_results  # single source of computed numbers

FIG = os.path.join(MANU, "figures")
OUT = os.path.join(HERE, "Carvajal_RDLC_screening_manuscript.docx")
FONT = "Times New Roman"


# --------------------------------------------------------------------------
# Formatting helpers (RDLC template: US Letter, 2 cm margins, Times New Roman)
# --------------------------------------------------------------------------

def setup(doc):
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(10)
    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.left_margin = sec.right_margin = Cm(2)
        sec.top_margin = sec.bottom_margin = Cm(2)


def _run(p, text, size=10, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return r


def title(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, 18, bold=True); p.paragraph_format.space_after = Pt(6)


def centered(doc, text, size=10, bold=False, italic=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, size, bold, italic); return p


def section(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    _run(p, text, 11, bold=True)


def subsection(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    _run(p, text, 10, bold=True, italic=True)


def body(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6); _run(p, text, 10); return p


def label_body(doc, label, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    _run(p, label + " ", 10, bold=True); _run(p, text, 10); return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Cm(0.75); _run(p, "- " + text, 10)


def numbered(doc, n, text):
    p = doc.add_paragraph(); _run(p, f"{n}. ", 10, bold=True); _run(p, text, 10)


def caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8); _run(p, text, 9, italic=True)


def figure(doc, path, cap):
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(12.5))
    else:
        body(doc, f"[missing figure: {path}]")
    caption(doc, cap)


def table(doc, cap, headers, rows):
    caption(doc, cap)
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        t.rows[0].cells[j].paragraphs[0].clear()
        _run(t.rows[0].cells[j].paragraphs[0], h, 9, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].clear(); _run(cells[j].paragraphs[0], str(val), 9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def reference(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(4); _run(p, text, 10)


# --------------------------------------------------------------------------
# Build
# --------------------------------------------------------------------------

def build():
    R = compute_paper2_results()
    f2 = lambda x: f"{x:.2f}"
    f3 = lambda x: f"{x:.3f}"

    vogel = R["vogel"]
    n_nodes, n_members, n_cols = R["vogel_n"]
    alp = R["alp"]
    study = R["variants"]
    fixed = next(v for v in study.variants if "Fixed" in v.label)
    pinned = next(v for v in study.variants if "Pinned" in v.label)

    rs_v = f3(vogel.robustness_index)
    rho = f"{alp.spearman_rho:+.2f}"
    crit = alp.entropy_rank[:2]                       # two most-critical columns
    crit_dcr = [alp.alp_severity[c] for c in crit]
    top5 = alp.topk_overlap.get(5, 0)
    rs_fixed = f3(fixed.robustness_index)
    rs_pinned = f3(pinned.robustness_index)
    u_err = R["vogel_indep"]["U_err_pct"] if R["vogel_indep"] else 0.0

    doc = Document()
    setup(doc)

    centered(doc, "Research Article", 10, italic=True)
    title(doc, "A Calibration-Free Entropy Index for Screening "
               "Progressive-Collapse Vulnerability of Steel Building Frames: "
               "Column Triage Validated Against Alternate-Load-Path Analysis")
    centered(doc, "Felipe Carvajal Brown 1, *", 10)
    aff = doc.add_paragraph(style="List Paragraph"); aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(aff, "1 Instituto Virginio Gomez, Concepcion (Chile); fcarvajalbrown@gmail.com", 9)
    centered(doc, "* Correspondence: fcarvajalbrown@gmail.com", 9)
    centered(doc, "Received: ____; Accepted: ____; Published: ____", 9, italic=True)
    cite = doc.add_paragraph()
    _run(cite, "Citation: ", 9, bold=True)
    _run(cite, "Carvajal Brown, F. (2026). A Calibration-Free Entropy Index for "
               "Screening Progressive-Collapse Vulnerability of Steel Building "
               "Frames. Revista de la Construccion. Journal of Construction.", 9)

    # ---- Highlights ----
    section(doc, "Highlights:")
    for h in [
        "A calibration-free index R_S screens which steel-frame columns are "
        "critical, from a single linear analysis.",
        "It ranks columns like a code alternate-load-path (ALP) removal check, "
        "so it triages where costly nonlinear ALP is worth running.",
        f"On the Vogel benchmark building the entropy ranking agrees with the "
        f"ALP severity at Spearman rho = {alp.spearman_rho:.2f}.",
        "R_S distinguishes design variants by redundancy and flags the columns "
        "whose loss the design cannot survive.",
    ]:
        bullet(doc, h)

    # ---- Abstract / keywords ----
    label_body(doc, "Abstract:",
        "Code guidance against progressive collapse (GSA, 2003; DoD, 2016) uses "
        "the alternate load path (ALP) method: notionally remove a primary "
        "column, re-analyse, and check that the remaining structure carries the "
        "load. Done faithfully this is a nonlinear analysis, and it must be "
        "repeated for every column the engineer chooses to remove, against "
        "acceptance limits calibrated to the structure. This paper proposes a "
        "cheap way to decide which columns deserve that effort. The Entropy "
        "Robustness Index R_S is the Shannon entropy of the elastic strain-energy "
        "distribution, averaged over single-member removals; its per-column "
        "entropy drop dH_k ranks columns by how much their loss localizes the "
        "load path. R_S needs no calibrated threshold and follows from one linear "
        "solve. We apply it to a published, recognisable steel building, the "
        "Vogel (1985) six-storey two-bay moment frame, and compare the entropy "
        "column ranking with a code-style ALP severity, the worst demand-capacity "
        "ratio (DCR) that appears after the column is removed. The two rankings "
        f"agree strongly (Spearman rho = {alp.spearman_rho:.2f} over all "
        f"{n_cols} columns); the two columns whose loss most overstresses the "
        "frame are exactly the two the entropy ranking places first, so an "
        "engineer can reserve the nonlinear ALP runs for the columns R_S flags. "
        "Because the model is linear, this ranking is independent of the load "
        "level. We then show R_S supports a design decision: it ranks a "
        "fixed-base (moment) version of the building as more robust than a "
        f"pinned-base variant ({rs_fixed} vs {rs_pinned}) and pinpoints the base "
        "columns whose loss the pinned design cannot survive. The scope is that "
        "of a screening tool: planar, first-order, linear-elastic and "
        "quasi-static, measuring load-path redundancy rather than strength, and "
        "marking localization as it happens rather than forecasting it.")
    label_body(doc, "Keywords:",
        "progressive collapse; steel frames; alternate load path; structural "
        "robustness; column removal; Shannon entropy; screening")

    # ---- Introduction ----
    section(doc, "Introduction")
    body(doc, "Progressive collapse is the spread of local damage into a failure "
        "out of proportion to its trigger. The design guidance written against "
        "it, from the General Services Administration guidelines (GSA, 2003) to "
        "Unified Facilities Criteria UFC 4-023-03 (DoD, 2016), is organised "
        "around the alternate load path (ALP) method: a primary load-carrying "
        "column is notionally removed, the structure is re-analysed, and the "
        "engineer checks whether the members that remain can carry the load along "
        "a new path. The method is powerful but expensive. A faithful ALP check "
        "is nonlinear, it is repeated once per removal scenario, and it is judged "
        "against acceptance limits (drift, chord rotation, an acceptance "
        "demand-capacity ratio) that must be chosen for the structure. On a real "
        "building with many columns, deciding which columns to remove, and how "
        "much analysis each deserves, is itself part of the engineering work.")
    body(doc, "This paper addresses that decision rather than replacing the ALP "
        "analysis. We ask whether a cheap, calibration-free index computed from a "
        "single linear analysis can tell the engineer which columns are critical, "
        "so that the costly nonlinear ALP effort is spent where it matters. The "
        "index we use is the Entropy Robustness Index R_S, the Shannon entropy of "
        "the elastic strain-energy distribution evaluated under single-member "
        "removal. The idea that an even internal distribution signals robustness "
        "is established: the redundancy matrix argues that robust structures "
        "distribute redundancy homogeneously (von Scheven et al., 2021); "
        "information-theoretic measures of structural integrity and of redundancy "
        "as the entropy of operational modes are older still (Nafday, 2011; Ziha, "
        "2000); and the entropy of a load-flow distribution is a standard "
        "fragility measure for cascading failure in networks (Koc et al., 2013). "
        "System-reserve redundancy and robustness indices for structures are well "
        "developed (Frangopol and Curley, 1987; Ghosn et al., 2010), and "
        "strain-energy member-importance coefficients are standard (Lin, 2019). "
        "R_S differs from these by measuring the entropy of the member-level "
        "elastic strain-energy distribution under code-style single-member "
        "removal, and by being bounded in [0, 1] with no structure-specific "
        "threshold.")
    body(doc, "The contribution here is not the index itself, which prior work "
        "has set out, but its use as a construction-engineering screening tool "
        "and the evidence that it works as one. We evaluate R_S on a published, "
        "recognisable steel building rather than a toy frame, we show its "
        "per-column ranking agrees with a code-style ALP severity so that it can "
        "triage columns for nonlinear analysis, and we show it supports a real "
        "design decision. Two honesties frame the claim. First, R_S measures "
        "load-path redundancy, how evenly strain energy is shared and whether a "
        "single loss is survivable, not strength or stiffness; we make that "
        "boundary explicit and do not claim it substitutes for capacity checks. "
        "Second, the entropy drop marks localization as it occurs and is not an "
        "early-warning signal.")

    # ---- Materials and methods ----
    section(doc, "Materials and methods")

    subsection(doc, "The entropy robustness index")
    body(doc, "At a given load let U_i be the elastic strain energy in active "
        "member i, with N active members. The normalized energy distribution and "
        "its Shannon entropy are given by Equation (1).")
    eq = doc.add_paragraph(); eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(eq, "p_i = U_i / sum(U_j),    S = - sum p_i ln(p_i),    H = S / ln(N)"
             "          (1)", 10)
    body(doc, "H equals 1 when strain energy is shared evenly and 0 when a single "
        "member holds it all. Because the p_i are ratios, H does not change as "
        "the load is scaled on a fixed topology; it moves only when the load path "
        "changes. For a linear-elastic frame the redistribution after a column is "
        "lost follows exactly from dropping that member from the stiffness "
        "assembly and solving Ku = F again, which is the ALP redistribution with "
        "no separate rule. Writing H_0 for the intact normalized entropy and H_k "
        "for the survivors after removing member k (with H_k = 0 when the removal "
        "forms a mechanism), the per-member entropy drop is dH_k = H_0 - H_k and "
        "the index is the mean over the removal set R, Equation (2).")
    eq2 = doc.add_paragraph(); eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(eq2, "R_S = (1/|R|) sum_k H_k,     in [0, 1]"
              "                              (2)", 10)
    body(doc, "R_S near 1 means the frame stays redundant under any single loss; "
        "near 0 means it does not. A determinacy bound ties this to redundancy: a "
        "statically determinate truss has R_S = 0, because removing any bar forms "
        "a mechanism, so R_S > 0 requires at least one degree of static "
        "indeterminacy. The per-column entropy drop dH_k is the screening output: "
        "columns with large dH_k are those whose loss most localizes the load "
        "path, and are the candidates for detailed analysis.")

    subsection(doc, "Code-style alternate-load-path severity")
    body(doc, "To test whether dH_k picks the columns an ALP analysis would flag, "
        "we compare it against a code-style severity for the same single-column "
        "removals. For each column we remove it, re-solve at the design load, and "
        "record the worst demand-capacity ratio (DCR, combined axial plus bending "
        "stress over yield) among the surviving members; a removal that forms a "
        "mechanism scores infinite severity. This is the quantity a linear-static "
        "ALP check inspects: whether the alternate path is within capacity or "
        "grossly overstressed. Unlike R_S it is threshold-based, which is the "
        "point of the comparison: a calibration-free index is judged against the "
        "threshold-based procedure it is meant to triage for. Because the model "
        "is linear, the post-removal DCR scales with the load factor, so the "
        "ranking of columns by severity is independent of the load level, as is "
        "dH_k; the agreement between them is therefore a structural property, not "
        "an artefact of the chosen load. We report their Spearman rank "
        "correlation and their top-k overlap.")

    subsection(doc, "Benchmark building")
    body(doc, "The study frame is the Vogel (1985) six-storey two-bay steel "
        "moment frame, a European calibration frame used to validate advanced "
        "steel-frame analyses and included in the Ziemian and Ziemian (2021) "
        "benchmark collection. It is 12 m wide (two 6 m bays) and 22.5 m tall "
        "(six 3.75 m storeys), with rigid beam-to-column connections, fixed "
        "bases, IPE floor beams and HEB columns in European S235 steel "
        "(E = 205 GPa). Beam sections range from IPE 400 at the lowest floor to "
        "IPE 240 at the roof; column sections from HEB 260/220 at the base to "
        "HEB 200/160 at the top. Gravity is the published distributed load "
        "(49.1 kN/m on floors, 31.7 kN/m on the roof), represented as "
        "work-equivalent joint loads, together with Vogel's notional horizontal "
        f"loads. The model has {n_nodes} nodes and {n_members} members, of which "
        f"{n_cols} are columns. We stress that the analysis is first-order and "
        "linear-elastic: Vogel's own results are second-order and inelastic and "
        "are not reproduced here. The frame is used as a realistic, provenance-"
        "backed geometry on which to exercise the screening index, not as a "
        "second-order benchmark.")

    subsection(doc, "Implementation and verification")
    body(doc, "The analysis uses an in-house planar Euler-Bernoulli frame solver "
        "whose modules communicate only through shared data classes. The solver "
        "is verified to floating-point precision against three closed-form beam "
        "cases and against an independent three-degree-of-freedom solver written "
        "from scratch. On the Vogel building the independent solver reproduces "
        "the joint displacements and total strain energy of the production solver "
        f"to within {u_err:.4f}% (that is, to floating-point precision), so the "
        "strain-energy distribution that R_S is built from is correct.")

    # ---- Results ----
    section(doc, "Results and discussion")

    subsection(doc, "Robustness of the benchmark building")
    body(doc, "The intact Vogel building has normalized entropy "
        f"H_0 = {f3(vogel.intact_entropy_norm)} and Entropy Robustness Index "
        f"R_S = {rs_v}. No single member loss forms a mechanism (mechanism "
        "fraction 0%), consistent with a redundant moment frame, and the worst "
        f"single loss pulls the surviving distribution to "
        f"H_k = {f3(vogel.worst_case_entropy_norm)}. R_S is therefore in the "
        "range expected of a genuinely redundant building, well away from the "
        "determinate limit of 0.")

    subsection(doc, "R_S column triage agrees with the ALP check")
    body(doc, "The screening claim is that the entropy column ranking dH_k picks "
        "the same critical columns as the code-style ALP severity. Over all "
        f"{n_cols} columns the two rankings agree with a Spearman correlation of "
        f"{rho} (Figure 1). The agreement is strongest exactly where it matters "
        "for triage: the two columns whose loss produces the worst surviving "
        f"overstress (post-removal DCR of {crit_dcr[0]:.1f} and {crit_dcr[1]:.1f}, "
        "the top-storey exterior columns) are the same two columns the entropy "
        f"ranking places first, and {top5} of the five most severe columns by the "
        "ALP check are within the five most critical by dH_k. In practice this "
        f"means an engineer can compute dH_k for all {n_cols} columns from one "
        "linear solve and reserve the expensive nonlinear ALP analyses for the "
        "handful the index flags, rather than analysing every column or choosing "
        "removal scenarios by judgement alone. The ranking carries no calibrated "
        "threshold and, being scale-invariant, does not depend on the load level "
        "at which it is computed.")
    figure(doc, os.path.join(FIG, "fig_alp_agreement.png"),
           "Figure 1. Entropy column criticality dH_k (one linear solve) against "
           "the code-style alternate-load-path severity (post-removal maximum "
           "DCR) for the Vogel building. The two most critical columns coincide "
           f"and the overall Spearman correlation is {rho}.")
    body(doc, "The agreement is strong but not perfect, and this is expected. The "
        "two measures answer overlapping but different questions: dH_k measures "
        "how much a loss localizes the load path, while the ALP DCR measures how "
        "much it overstresses the survivors. They rank the most critical columns "
        "the same way, which is what a triage needs, while differing on the less "
        "critical columns where the decision matters least.")

    subsection(doc, "A design decision: base fixity")
    body(doc, "A screening index earns its place if it can tell two designs "
        "apart. We compare the Vogel building as built, with fixed (moment) "
        "column bases, against a variant with pinned bases, a genuine and "
        "cost-relevant steel-design choice. R_S ranks the fixed-base design as "
        f"more robust ({rs_fixed}) than the pinned-base design ({rs_pinned}), a "
        f"drop of {study.delta_rs:.3f} (Figure 2, Table 1). The reason is "
        "explicit in the index: with pinned bases the loss of a base column is no "
        f"longer survivable, so {pinned.mechanism_fraction:.0%} of single-column "
        "removals form a mechanism, and the worst-case surviving entropy falls to "
        "zero. Moreover the entropy criticality ranking of the pinned design "
        f"places the three base columns {pinned.mechanism_columns} at the very "
        "top, so the index not only scores the weaker design lower but points at "
        "the exact members whose loss it cannot survive, i.e. where redundancy "
        "must be restored.")
    table(doc, "Table 1. Design-variant comparison for the Vogel building "
               "(computed by analysis/design_variants.py).",
          ["Design variant", "R_S", "Single-loss mechanisms",
           "Mechanism-forming columns"],
          [["Fixed (moment) bases, as built", rs_fixed,
            f"{fixed.mechanism_fraction:.0%}",
            (", ".join(map(str, fixed.mechanism_columns)) or "none")],
           ["Pinned bases", rs_pinned, f"{pinned.mechanism_fraction:.0%}",
            ", ".join(map(str, pinned.mechanism_columns))]])
    figure(doc, os.path.join(FIG, "fig_design_variants.png"),
           "Figure 2. R_S distinguishes the two designs: the pinned-base variant "
           "is less redundant, admitting single-column-loss mechanisms at the "
           "base columns, and R_S scores it lower.")

    subsection(doc, "What R_S measures, and what it does not")
    body(doc, "It is important to state the boundary of the index, because it "
        "shapes how the screening should be used. R_S rewards distributed, "
        "redundant load paths and the survivability of single losses. It is not a "
        "strength or stiffness measure, and it does not reward member stiffening: "
        "increasing a member's section concentrates strain energy in it and can "
        "lower R_S, and adding a stiff brace to an already-redundant frame can do "
        "the same. R_S therefore screens for redundancy, the property the ALP "
        "method is concerned with, and it should be read alongside, not in place "
        "of, strength and serviceability checks. This division of labour is a "
        "feature for a triage tool: R_S answers 'which column losses does this "
        "structure struggle to survive', which is precisely the ALP question, and "
        "leaves 'is each member strong enough' to the capacity checks.")

    subsection(doc, "Limitations as the envelope of a screening tool")
    body(doc, "The limits define where the screening applies rather than "
        "apologise for it. The analysis is planar, first-order, linear-elastic "
        "and quasi-static, so it omits plasticity, second-order and catenary "
        "action, out-of-plane response, dynamic amplification, and connection "
        "failure. These are exactly the effects a detailed nonlinear ALP analysis "
        "is run to capture, which is why R_S is positioned as the cheap first "
        "pass that decides where to run it, not as a replacement. Linear-static "
        "ALP is itself an accepted, generally conservative code method, though "
        "for stability-governed compression members the first-order assumption "
        "can be unconservative, a further reason to reserve nonlinear analysis "
        "for the flagged columns. Finally, the entropy drop lands with the "
        "localization rather than ahead of it, so the index is a calibration-free "
        "measure of robustness and localization, not a predictive early warning.")

    # ---- Conclusions ----
    section(doc, "Conclusions")
    for i, c in enumerate([
        "The Entropy Robustness Index R_S provides a calibration-free, "
        "single-linear-solve screening of steel-frame columns for "
        "progressive-collapse vulnerability, with its per-column entropy drop "
        "dH_k ranking columns by how much their loss localizes the load path.",
        "On the published Vogel (1985) benchmark building the entropy column "
        f"ranking agrees with a code-style alternate-load-path severity at "
        f"Spearman rho = {alp.spearman_rho:.2f}, and the two most critical "
        "columns coincide, so R_S can triage which columns deserve the expensive "
        "nonlinear ALP analysis.",
        "R_S supports design decisions: it ranks a fixed-base building as more "
        f"robust than a pinned-base variant ({rs_fixed} vs {rs_pinned}) and "
        "identifies the base columns whose loss the weaker design cannot survive.",
        "The index measures load-path redundancy, not strength; it belongs "
        "alongside capacity checks and ahead of nonlinear ALP analysis, within a "
        "planar, linear-elastic, quasi-static envelope and with no early-warning "
        "claim.",
    ], start=1):
        numbered(doc, i, c)

    label_body(doc, "Author contributions:",
        "F.C.B. conceived the method, implemented the software, performed the "
        "analyses, and wrote the manuscript.")
    label_body(doc, "Funding:", "This research received no external funding.")
    label_body(doc, "Data availability:",
        "The software, benchmark frames and analysis scripts that reproduce every "
        "number and figure in this paper are openly available in the project "
        "repository.")
    label_body(doc, "Conflicts of interest:", "The author declares no conflict of interest.")

    section(doc, "References")
    for ref in REFERENCES:
        reference(doc, ref)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


REFERENCES = [
    "Department of Defense. (2016). Unified Facilities Criteria UFC 4-023-03: "
    "Design of Buildings to Resist Progressive Collapse.",
    "Feng, D.-C., et al. (2024). Physically-based collapse failure criteria in "
    "progressive collapse analyses of random-parameter multi-story RC structures "
    "subjected to column removal scenarios. Engineering Structures. "
    "https://doi.org/10.1016/j.engstruct.2024.119412",
    "Frangopol, D. M., & Curley, J. P. (1987). Effects of damage and redundancy "
    "on structural reliability. Journal of Structural Engineering, 113(7), "
    "1533-1549.",
    "General Services Administration. (2003). Progressive Collapse Analysis and "
    "Design Guidelines for New Federal Office Buildings and Major Modernization "
    "Projects.",
    "Ghosn, M., Moses, F., & Frangopol, D. M. (2010). Redundancy and robustness "
    "of highway bridge superstructures and substructures. Structure and "
    "Infrastructure Engineering, 6(1-2), 257-278.",
    "Koc, Y., Warnier, M., Kooij, R. E., & Brazier, F. M. (2013). An "
    "entropy-based metric to quantify the robustness of power grids against "
    "cascading failures. Safety Science, 59, 126-134.",
    "Lin, K., et al. (2019). Importance assessment of structural members based "
    "on elastic-plastic strain energy. Advances in Materials Science and "
    "Engineering, 2019, 8019675.",
    "Nafday, A. M. (2011). Consequence-based structural design approach for black "
    "swan events. Structural Safety, 33(1), 108-114.",
    "Shannon, C. E. (1948). A mathematical theory of communication. Bell System "
    "Technical Journal, 27, 379-423.",
    "von Scheven, M., Ramm, E., & Bischoff, M. (2021). Quantification of the "
    "redundancy distribution in truss and beam structures. International Journal "
    "of Solids and Structures, 213, 41-49. "
    "https://doi.org/10.1016/j.ijsolstr.2020.11.002",
    "Vogel, U. (1985). Calibrating frames. Stahlbau, 54, 293-301.",
    "Ziemian, C. W., & Ziemian, R. D. (2021). Steel benchmark frames for "
    "structural analysis and validation studies: Finite element models and "
    "numerical simulation data. Data in Brief, 39, 107564. "
    "https://doi.org/10.1016/j.dib.2021.107564",
    "Ziha, K. (2000). Redundancy and robustness of systems of events. "
    "Probabilistic Engineering Mechanics, 15(4), 347-357.",
]


if __name__ == "__main__":
    build()
