"""Generate the Ingeniare cover letter (carta de presentacion).

Follows Modelo_Carta_de_presentacion_Articulos_Ingeniare.doc. Single author
(F. Carvajal Brown). Placeholders marked [COMPLETAR: ...] must be filled by
the author: the city/date, the five suggested reviewers, and the signature.
Nothing about real people (reviewers) is invented here.
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Carvajal_Carta_Presentacion_Ingeniare.docx")
TITLE = ("Entropía de Shannon como indicador de colapso progresivo en marcos "
         "estructurales: un simulador basado en MEF")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(12)


def para(text="", align=None, bold=False, space_after=10):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold
    return p


para("CARTA DE PRESENTACIÓN DE ARTÍCULO Y CESIÓN DE DERECHOS DE AUTOR",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

para("Ñuñoa, Santiago de Chile, 18 de junio de 2026",
     align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

para("Señor", space_after=0)
para("Kristopher Chandía Valenzuela", space_after=0)
para("Editor", space_after=0)
para("Ingeniare. Revista chilena de ingeniería", space_after=0)
para("Presente", space_after=18)

para("Remito el manuscrito titulado «" + TITLE + "» para que sea sometido al "
     "proceso de evaluación y eventual publicación en Ingeniare. Revista "
     "chilena de ingeniería.")

para("El autor declara:", space_after=4)
for d in [
    "Que es un trabajo original.",
    "Que no ha sido previamente publicado en otro medio.",
    "Que no ha sido remitido paralelamente a otro medio de publicación.",
    "Que ha contribuido intelectualmente en su elaboración.",
    "Que ha leído y aprobado la versión final del manuscrito remitido.",
    "Que, en caso de ser aprobado y publicado el artículo, cede todos los "
    "derechos de publicación a Ingeniare. Revista chilena de ingeniería.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(d)

para("")
para("Cabe señalar que el autor de correspondencia es: Felipe Carvajal Brown "
     "(fcarvajalbrown@gmail.com).")

para("En caso de que el artículo sea aprobado para publicación, a través de la "
     "presente, Ingeniare. Revista chilena de ingeniería, asume los derechos "
     "exclusivos para editar, publicar, reproducir, distribuir copias (formato "
     "impreso y/o electrónico) e incluir el artículo en índices nacionales e "
     "internacionales o bases de datos.")

para("Recomiendo como posibles evaluadores del artículo a los siguientes "
     "expertos en el tema (no vinculados a la institución del autor y sin "
     "conflicto de interés). Para cada uno se indica nombre completo, "
     "filiación, línea de investigación, ORCID y correo electrónico:",
     space_after=6)

REVIEWERS = [
    dict(nombre="Dr. Rodrigo Astroza Eulufí",
         filiacion="Facultad de Ingeniería y Ciencias Aplicadas, "
                   "Universidad de los Andes, Santiago, Chile",
         linea="Monitoreo de salud estructural, actualización de modelos de "
               "elementos finitos e identificación de daño, dinámica "
               "estructural",
         orcid="0000-0003-0711-1259",
         email="rastroza@uandes.cl"),
    dict(nombre="Dr. Rubén Boroschek Krauskopf",
         filiacion="Departamento de Ingeniería Civil, Universidad de Chile, "
                   "Santiago, Chile",
         linea="Monitoreo de salud estructural, dinámica estructural e "
               "ingeniería sísmica",
         orcid="0000-0003-2253-2334",
         email="rborosch@uchile.cl"),
    dict(nombre="Dr. Gastón Fermandois Cornejo",
         filiacion="Departamento de Obras Civiles, Universidad Técnica "
                   "Federico Santa María, Valparaíso, Chile",
         linea="Control estructural, monitoreo de salud estructural y "
               "detección de daño, simulación híbrida en tiempo real",
         orcid="0000-0001-7320-7991",
         email="gaston.fermandois@usm.cl"),
    dict(nombre="Dr. Héctor Jensen",
         filiacion="Departamento de Obras Civiles, Universidad Técnica "
                   "Federico Santa María, Valparaíso, Chile",
         linea="Mecánica estructural estocástica, confiabilidad y "
               "optimización bajo incertidumbre, análisis de riesgo",
         orcid="[completar]",
         email="hector.jensen@usm.cl"),
    dict(nombre="Dr. Juan Felipe Beltrán Morales",
         filiacion="Departamento de Ingeniería Civil, Universidad de Chile, "
                   "Santiago, Chile",
         linea="Mecánica estructural computacional, estructuras reticuladas "
               "y de cables",
         orcid="0000-0002-2054-8079",
         email="jbeltran@ing.uchile.cl"),
]
for n, rv in enumerate(REVIEWERS, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(f"{n}. {rv['nombre']}. ").bold = True
    p.add_run(f"{rv['filiacion']}. Línea de investigación: {rv['linea']}. "
              f"ORCID: {rv['orcid']}. E-mail: {rv['email']}.")

para("")
para("Se adjuntan los siguientes documentos: manuscrito en formato Word "
     "(.docx) y PDF. Las figuras están incluidas dentro del manuscrito.")

para("Atentamente,", space_after=24)
para("[COMPLETAR: firma]", space_after=0)
para("Felipe Carvajal Brown", space_after=0)
para("fcarvajalbrown@gmail.com", space_after=0)
para("ORCID: 0000-0002-8300-7587", space_after=0)

doc.save(OUT)
print("saved ->", OUT)
