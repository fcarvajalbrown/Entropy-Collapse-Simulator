"""Apply the Ingeniare editorial fixes to paper_ingeniare.docx.

Round 1 (Carina Garcia, Revista Ingeniare):
  - Cite [3]-[5]  -> citation-format fix (in-text superscript parentheses).
  - Some equations are images -> convert to editable native Word equations.
  - Full journal names + format (quotes, commas, periods) in references.
Plus template-compliance: bold Spanish title and section headings, keyword
tweak, centered RESUMEN/ABSTRACT labels.

Round 2 (Carina Garcia):
  - Conclusions must not be numbered  -> strip "N) " prefixes.
  - Cite every equation in the text as "ecuacion (N)"  -> add refs to 2,3,4.
  - References not tabular  -> remove hanging indent.
  - Figure 2 has a broken symbol  -> repair detached acute accent in the
    "S / S_max" y-axis label (also Figure 5, same defect). Plot data is NOT
    touched: only the label glyphs are repainted on the existing PNG.

Scope: FORMAT ONLY. No scientific content is altered. The equations reproduce
the exact content currently in the submitted paper.

Pipeline for equations: LaTeX -> MathML (latex2mathml) -> OMML (Office XSLT).
Output is written to a NEW file; the original is left untouched.
"""
import copy
import io
import os
import re
import shutil
import zipfile

import lxml.etree as ET
import latex2mathml.converter as conv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "paper_ingeniare.docx")
DST = os.path.join(HERE, "Carvajal_Ingeniare_manuscript_v3.docx")
XSLT = r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL"
FONT = (r"C:\Users\Beetlejuice\AppData\Local\Python\pythoncore-3.14-64\Lib"
        r"\site-packages\matplotlib\mpl-data\fonts\ttf\DejaVuSans.ttf")

# Figure media files (in document order) whose "S / S_max" y-label is broken.
FIGURE_LABEL_FIX = [
    "word/media/558cd470ee9f037847bef7b09eacea2133e1e570.png",  # Figura 2
    "word/media/67a0b4d7c9067bcdc2e741e6010fab971c80f2e5.png",  # Figura 5
]

_transform = ET.XSLT(ET.parse(XSLT))


def _label_image(width_px):
    """Render a clean 'S / S_max' (max as subscript) rotated for a y-axis."""
    main = ImageFont.truetype(FONT, 60)
    sub = ImageFont.truetype(FONT, 37)
    s1, s2 = "S / S", "máx"
    probe = ImageDraw.Draw(Image.new("RGBA", (10, 10)))
    w1 = probe.textlength(s1, font=main)
    asc, desc = main.getmetrics()
    img = Image.new("RGBA",
                    (int(w1 + probe.textlength(s2, font=sub) + 10),
                     asc + desc + 60), (255, 255, 255, 0))
    d = ImageDraw.Draw(img)
    d.text((0, 0), s1, font=main, fill=(0, 0, 0, 255))
    d.text((w1, 19), s2, font=sub, fill=(0, 0, 0, 255))
    lbl = img.crop(img.getbbox()).rotate(90, expand=True)
    scale = width_px / lbl.width
    return lbl.resize((width_px, int(lbl.height * scale)), Image.LANCZOS)


def fix_figure_label(png_bytes):
    """Repaint the broken y-axis label on a figure PNG (data untouched)."""
    im = Image.open(io.BytesIO(png_bytes)).convert("RGB")
    px = im.load()
    xs, ys = [], []
    for y in range(im.height):
        for x in range(95):           # left strip = rotated axis label only
            r, g, b = px[x, y]
            if r < 120 and g < 120 and b < 120:
                xs.append(x)
                ys.append(y)
    if not xs:
        return png_bytes              # nothing detected; leave as-is
    minx, maxx, miny, maxy = min(xs), max(xs), min(ys), max(ys)
    ImageDraw.Draw(im).rectangle(
        [minx - 6, miny - 16, maxx + 6, maxy + 16], fill=(255, 255, 255))
    lbl = _label_image((maxx - minx) + 1)
    cx, cy = (minx + maxx) // 2, (miny + maxy) // 2
    im.paste(lbl, (cx - lbl.width // 2, cy - lbl.height // 2), lbl)
    out = io.BytesIO()
    im.save(out, format="PNG")
    return out.getvalue()


def replace_zip_entries(docx_path, transforms):
    """Rewrite a .docx, applying transforms[arcname](bytes)->bytes."""
    buf = io.BytesIO()
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename in transforms:
                    data = transforms[item.filename](data)
                zout.writestr(item, data)
    with open(docx_path, "wb") as f:
        f.write(buf.getvalue())


_M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def _nest_nary_operands(omath):
    """Move a summand that follows an n-ary (empty operand) into the n-ary.

    latex2mathml + MML2OMML render "\\sum_j X" as an n-ary with an EMPTY
    <m:e> followed by X as siblings, which Word draws as a stray empty box.
    For our equations everything after the n-ary is its summand, so move it in.
    """
    for nary in list(omath.iter(_M + "nary")):
        e = nary.find(_M + "e")
        if e is None or len(e) > 0:
            continue
        parent = nary.getparent()
        after = False
        movers = []
        for child in list(parent):
            if child is nary:
                after = True
                continue
            if after:
                movers.append(child)
        for m in movers:
            parent.remove(m)
            e.append(m)
    return omath


def latex_to_omath(latex):
    """Return an lxml <m:oMath> element for the given LaTeX string."""
    mml = conv.convert(latex)
    dom = ET.fromstring(mml.encode("utf-8"))
    omml = _transform(dom)
    # omml is an _XSLTResultTree whose root is <m:oMath>
    return _nest_nary_operands(copy.deepcopy(omml.getroot()))


# --- Equation paragraph index -> (LaTeX, trailing plain text incl. number) ---
EQUATIONS = {
    29: (r"U_i = \frac{1}{2}\cdot u_i^{loc}\cdot k_i^{loc}\cdot u_i^{loc}",
         "  (1)"),
    31: (r"S = -\sum_{i=1}^{N} p_i \ln(p_i)",
         "  (2)"),
    39: (r"\sigma_{max} = \frac{|N|}{A} + \frac{|M_{max}|\cdot c}{I}",
         "  (3)"),
    41: (r"\frac{dU_i}{dt} = \sum_j \alpha_{ij}\,(U_j - U_i)",
         "  (4)"),
}


def replace_equation(paragraph, latex, trailer):
    p = paragraph._p
    # remove existing run elements only (keep w:pPr)
    for r in p.findall(qn("w:r")):
        p.remove(r)
    # append the editable equation
    p.append(latex_to_omath(latex))
    # append the equation number / annotation as a normal run
    run = paragraph.add_run(trailer)
    return run


# --- References: list of (text, italic) segments per reference number ---
Q1, Q2 = "“", "”"  # curly quotes
F, I = False, True
REFERENCES = {
    1: [("[1] U. Starossek, ", F),
        ("Progressive Collapse of Structures", I),
        (". London, U.K.: Thomas Telford, 2009.", F)],
    2: [("[2] U.S. General Services Administration (GSA), ", F),
        ("Progressive Collapse Analysis and Design Guidelines for New Federal "
         "Office Buildings and Major Modernization Projects", I),
        (". Washington, DC, USA, 2003.", F)],
    3: [("[3] D. Feng, G. Wu, and R. Lu, " + Q1 +
         "Physically-based collapse failure criteria in progressive collapse "
         "analyses of random-parameter multi-story RC structures," + Q2 + " ", F),
        ("Journal of Building Engineering", I),
        (", vol. 91, 2024, doi: 10.1016/j.jobe.2024.019412.", F)],
    4: [("[4] A. Moreno-Gomez et al., " + Q1 +
         "EMD-Shannon entropy-based methodology to detect incipient damages in "
         "a truss structure," + Q2 + " ", F),
        ("Applied Sciences", I),
        (", vol. 8, no. 11, p. 2068, 2018, doi: 10.3390/app8112068.", F)],
    5: [("[5] T.-K. Lin and A. G. Laínez, " + Q1 +
         "Entropy-based structural health monitoring system for damage "
         "detection in multi-bay three-dimensional structures," + Q2 + " ", F),
        ("Entropy", I),
        (", vol. 20, no. 1, p. 49, 2018, doi: 10.3390/e20010049.", F)],
    6: [("[6] C. E. Shannon, " + Q1 +
         "A mathematical theory of communication," + Q2 + " ", F),
        ("Bell System Technical Journal", I),
        (", vol. 27, no. 3, pp. 379-423, 1948, "
         "doi: 10.1002/j.1538-7305.1948.tb01338.x.", F)],
    7: [("[7] K. J. Bathe, ", F),
        ("Finite Element Procedures", I),
        (", 2nd ed. Watertown, MA, USA: K. J. Bathe, 2014.", F)],
    8: [("[8] S. P. Timoshenko and J. M. Gere, ", F),
        ("Theory of Elastic Stability", I),
        (", 2nd ed. New York, NY, USA: McGraw-Hill, 1961.", F)],
    9: [("[9] J. P. Amezquita-Sanchez et al., " + Q1 +
         "Vibration control of smart civil structures with magneto-rheological "
         "dampers: A review," + Q2 + " ", F),
        ("Engineering Structures", I),
        (", vol. 208, 2020, doi: 10.1016/j.engstruct.2020.110058.", F)],
    10: [("[10] F. Carvajal Brown, " + Q1 + "Entropy Collapse Simulator," + Q2 +
          " GitHub, 2025. [En línea]. Disponible en: "
          "https://github.com/fcarvajalbrown/entropy-collapse-simulator. "
          "Accedido: mar. 2026.", F)],
}


def rebuild_reference(paragraph, segments):
    p = paragraph._p
    for r in p.findall(qn("w:r")):
        p.remove(r)
    for text, italic in segments:
        run = paragraph.add_run(text)
        run.italic = italic


CITE_RE = re.compile(r"^\(\d+(?:\s*,\s*\d+)*\)$")


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    paras = doc.paragraphs

    # 0. Spanish title -> bold (template requires "centrado en negrita")
    for r in paras[0].runs:
        r.bold = True
    print("spanish title bolded")

    # 0b. Level-1 section headings (16pt, centered, ALL CAPS) -> bold
    # Template: "Titulos: centrados en negrita con mayusculas, Calibri 16 pt."
    nh = 0
    for p in paras:
        t = p.text.strip()
        if not t or t != t.upper():
            continue
        sizes = {round(r.font.size.pt, 1) for r in p.runs
                 if r.font.size and r.text.strip()}
        if sizes == {16.0} and p.alignment is not None \
                and "CENTER" in str(p.alignment):
            for r in p.runs:
                r.bold = True
            nh += 1
    print("section headings bolded:", nh)

    # 1. Equations -> editable
    for idx, (latex, trailer) in EQUATIONS.items():
        replace_equation(paras[idx], latex, trailer)
    print("equations converted:", sorted(EQUATIONS))

    # find references section start
    refstart = next(i for i, p in enumerate(paras)
                    if p.text.strip().upper().startswith("REFERENCIA"))

    # 2/3. Reference list -> [N], italics, curly quotes, IEEE punctuation
    num_re = re.compile(r"^\((\d+)\)")
    done = set()
    for i in range(refstart + 1, len(paras)):
        m = num_re.match(paras[i].text.strip())
        if not m:
            continue
        n = int(m.group(1))
        if n in REFERENCES:
            rebuild_reference(paras[i], REFERENCES[n])
            done.add(n)
    print("references reformatted:", sorted(done))

    # 4. Keywords: reduce overlap with title, drop generic "metodo".
    #    Replacement terms are all genuinely discussed in the paper
    #    (structural redundancy, Gini index) - no invented content.
    KEYWORDS = {
        "Palabras clave: colapso progresivo, redundancia estructural, "
        "energía de deformación, índice de Gini, Euler-Bernoulli.":
            ("Palabras clave:", "palabras"),
        "Keywords: progressive collapse, structural redundancy, strain "
        "energy, Gini index, Euler-Bernoulli.":
            ("Keywords:", "keywords"),
    }
    for new_text, (_, marker) in KEYWORDS.items():
        for p in paras:
            low = p.text.strip().lower()
            if low.startswith(marker):
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
                p.add_run(new_text)
                break
    print("keywords updated")

    # 5. Center the RESUMEN: / ABSTRACT: labels (match house layout)
    for p in paras:
        if p.text.strip().upper() in ("RESUMEN:", "ABSTRACT:"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print("resumen/abstract labels centered")

    # 6. Cite every equation in the text (eq 1 already cited at "ecuacion (1)").
    #    Minimal, faithful edits to the sentences that introduce eqs 2, 3, 4.
    EQ_CITE = [
        ("Esta expresión alcanza su máximo S",
         "La ecuación (2) alcanza su máximo S"),
        ("se evalúa mediante tensión combinada axial y de flexión:",
         "se evalúa mediante la tensión combinada axial y de flexión, "
         "ecuación (3):"),
        ("la energía se redistribuye mediante:",
         "la energía se redistribuye mediante la ecuación (4):"),
    ]
    for old, new in EQ_CITE:
        for p in paras:
            for r in p.runs:
                if old in r.text:
                    r.text = r.text.replace(old, new)
    print("equation citations added (2,3,4)")

    # 6b. Validation formulas (delta = PL^3/48EI, U = P^2 L^3/96EI): the
    #     exponents are plain digits in their own runs; make them superscript.
    for p in paras:
        if "48EI" in p.text and "96EI" in p.text:
            for r in p.runs:
                if r.text in ("2", "3"):
                    r.font.superscript = True
            break
    print("validation formula exponents superscripted")

    # 7. Conclusions must not be numbered: strip leading "N) ".
    ci = next(i for i, p in enumerate(paras)
              if p.text.strip().upper() == "CONCLUSIONES")
    nconc = 0
    for p in paras[ci + 1:]:
        if p.runs and re.match(r"^\d\)\s", p.runs[0].text):
            p.runs[0].text = re.sub(r"^\d\)\s*", "", p.runs[0].text)
            nconc += 1
        elif p.text.strip().upper() in ("REFERENCIAS", "AGRADECIMIENTOS"):
            break
    print("conclusion items de-numbered:", nconc)

    # 8. References not tabular: remove hanging indent from each entry.
    refstart = next(i for i, p in enumerate(paras)
                    if p.text.strip().upper().startswith("REFERENCIA"))
    nind = 0
    for p in paras[refstart + 1:]:
        if re.match(r"^\[\d+\]", p.text.strip()):
            pPr = p._p.find(qn("w:pPr"))
            if pPr is not None:
                ind = pPr.find(qn("w:ind"))
                if ind is not None:
                    pPr.remove(ind)
                    nind += 1
    print("reference entries de-indented:", nind)

    # 1b. In-text citations -> superscript (body only, before references)
    sup = 0
    for i in range(refstart):
        if i in EQUATIONS:
            continue
        for r in paras[i].runs:
            if CITE_RE.match(r.text.strip()):
                r.font.superscript = True
                sup += 1
    print("citations superscripted:", sup)

    doc.save(DST)

    # 9. Repair the broken y-axis label on Figures 2 and 5 (PNG, data intact).
    replace_zip_entries(
        DST, {name: fix_figure_label for name in FIGURE_LABEL_FIX})
    print("figure labels repaired:", len(FIGURE_LABEL_FIX))

    print("saved ->", DST)


if __name__ == "__main__":
    main()
