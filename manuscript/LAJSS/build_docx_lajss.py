"""
manuscript/build_docx_lajss.py
==============================
Render the manuscript for the Latin American Journal of Solids and Structures
(LAJSS) by reusing the official "LAJSS - Template.docx" found in the repository
root. The template is opened as the base document so that all of its named
paragraph styles (ArticleTitle, Authors, Affiliations, Abstract, Heading1,
Heading2, Equation, FigureCaption, TableTitle, dataStatement,
fnAuthor-notesContributions, Reference, ...) are available; the placeholder body
is then cleared and refilled with our content.

Every reported number is COMPUTED at build time from the analysis modules and
the validation harness (it reuses compute_results() from build_docx.py), so the
LAJSS manuscript cannot drift out of sync with the code or with the RDLC copy.

LAJSS house rules honoured here:
  - Abstract body <= 1,140 characters with spaces (checked, raises if exceeded).
  - Numbered upper-case Heading1, lower-case Heading2.
  - In-text citation style Author (year); LAJSS reference list format.
  - CRediT author-contributions statement and a Data Availability statement.
  - ORCID is compulsory for every author (placeholder flagged below).

    python manuscript/generate_figures.py             # produce figures first
    python manuscript/LAJSS/build_docx_lajss.py        # -> manuscript/LAJSS/Carvajal_LAJSS_manuscript.docx
"""

import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
MANU = os.path.abspath(os.path.join(HERE, ".."))   # manuscript/ (shared results.py + figures/)
sys.path.insert(0, MANU)

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from results import compute_results  # single source of computed numbers

TEMPLATE = os.path.join(HERE, "LAJSS - Template.docx")  # drop the journal template here to rebuild
FIG = os.path.join(MANU, "figures")
OUT = os.path.join(HERE, "Carvajal_LAJSS_manuscript.docx")

ORCID = "https://orcid.org/0000-0002-8300-7587"
ABSTRACT_MAX_CHARS = 1140  # LAJSS limit: body of the abstract, characters with spaces


# --------------------------------------------------------------------------
# Template helpers
# --------------------------------------------------------------------------

def clear_body(doc):
    """Remove every placeholder paragraph and table, keep the section properties."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag
        if tag.endswith("}p") or tag.endswith("}tbl"):
            body.remove(child)


def has_style(doc, name):
    return name in (s.name for s in doc.styles)


def P(doc, text, style):
    """Add a paragraph in a named template style (fallback to Normal if absent)."""
    return doc.add_paragraph(text, style=style if has_style(doc, style) else None)


def figure(doc, path, cap):
    p = doc.add_paragraph(style="Figure" if has_style(doc, "Figure") else None)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Cm(13.0))
    else:
        p.add_run(f"[missing figure: {os.path.basename(path)}]")
    P(doc, cap, "FigureCaption").alignment = WD_ALIGN_PARAGRAPH.CENTER


def table(doc, title_text, headers, rows):
    P(doc, title_text, "TableTitle")
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].clear()
            cells[j].paragraphs[0].add_run(str(val)).font.size = Pt(9)
    doc.add_paragraph()


def equation(doc, text, number):
    """Equation paragraph with a right-ragged number, in the template Equation style."""
    p = P(doc, "", "Equation")
    p.add_run(text + "\t(" + str(number) + ")")
    return p


# --------------------------------------------------------------------------
# Build
# --------------------------------------------------------------------------

def build():
    R = compute_results()
    f2 = lambda x: f"{x:.2f}"
    f3 = lambda x: f"{x:.3f}"

    rs_beam = f2(R["beam"].robustness_index)
    rs_frame = f2(R["frame"].robustness_index)
    rs_truss = f2(R["truss"].robustness_index)
    rho_f = f"{R['imp_frame'].spearman_rho:+.2f}"
    rho_t = f"{R['imp_truss'].spearman_rho:+.2f}"

    if not os.path.exists(TEMPLATE):
        raise SystemExit(f"LAJSS template not found at {TEMPLATE}")
    doc = Document(TEMPLATE)
    clear_body(doc)

    # ---- Title block ----
    P(doc, "An Entropy Robustness Index for Planar Steel Frames: Formulation, "
           "Verified Implementation, and Comparison with Standard Collapse Criteria",
      "ArticleTitle")
    P(doc, f"Felipe Carvajal Brown a,* {ORCID}", "Authors")
    P(doc, "a Instituto Virginio Gomez, Concepcion, Chile. "
           "Email: fcarvajalbrown@gmail.com", "Affiliations")
    P(doc, "* Corresponding author", "Correspondence")

    # ---- Abstract ----
    P(doc, "Abstract", "AbstractTitle")
    abstract = (
        "Structural robustness rarely reduces to a single number. Engineers judge "
        "it with displacement, resistance, or energy criteria, and each needs a "
        "limit calibrated to the structure at hand. This paper proposes the "
        "Entropy Robustness Index R_S, the Shannon entropy of the elastic "
        "strain-energy distribution averaged over the notional member removals of "
        "the alternate-load-path method. R_S sits near 1 when a frame keeps "
        "sharing energy evenly after any single loss, and falls toward 0 when a "
        "loss localizes the energy or forms a mechanism. Implemented on a "
        "first-order linear-elastic planar Euler-Bernoulli solver verified to "
        f"floating-point precision, R_S reads {rs_beam}, {rs_frame}, and "
        f"{rs_truss} for a non-redundant beam, a redundant moment frame, and a "
        "redundant truss. It rises monotonically with the degree of static "
        "indeterminacy, and its member ranking is not a re-encoding of "
        "compliance importance. On one incremental analysis the entropy criterion "
        "reaches the collapse onset at the same load as the resistance and energy "
        "criteria, with no calibrated threshold."
    )
    n = len(abstract)
    if n > ABSTRACT_MAX_CHARS:
        raise SystemExit(f"Abstract is {n} chars > LAJSS limit {ABSTRACT_MAX_CHARS}.")
    print(f"Abstract length: {n} / {ABSTRACT_MAX_CHARS} chars with spaces")
    P(doc, abstract, "Abstract")

    P(doc, "Keywords", "KeywordsTitle")
    P(doc, "structural robustness; progressive collapse; alternate load path; "
           "Shannon entropy; redundancy", "Keywords")

    # ---- 1 Introduction ----
    P(doc, "1 INTRODUCTION", "Heading1")
    P(doc, "Progressive collapse is the spread of local damage into a failure out "
        "of proportion to what caused it. Guidance against it, from the GSA (2003) "
        "guidelines to Unified Facilities Criteria UFC 4-023-03 (Department of "
        "Defense, 2016), is built around the alternate load path (ALP) method: "
        "remove a primary load-carrying member, re-analyse the structure, and "
        "check whether the members that remain can carry the load along a new "
        "path. What that check should measure is less settled. Reviewing collapse "
        "criteria for reinforced-concrete frames under column removal, Feng et al. "
        "(2024) compare displacement, resistance, and energy criteria and find no "
        "consensus among them. Each one rests on a calibrated limit, whether a "
        "drift or chord-rotation limit, an acceptance demand-capacity ratio, or an "
        "energy capacity.", "Paragraph")
    P(doc, "The idea that an even internal distribution signals robustness is not "
        "new, and it has been expressed in information-theoretic terms before. The "
        "redundancy matrix describes how static indeterminacy is spread across "
        "members and argues that robust structures distribute redundancy "
        "homogeneously (von Scheven et al., 2021); like R_S it is load-independent "
        "and threshold-free. Nafday (2011) casts structural integrity in "
        "information-theoretic terms, and event-oriented system analysis defines "
        "redundancy directly as the Shannon entropy of operational modes (Ziha, "
        "2000). In network science the Shannon entropy of a load-flow "
        "distribution is an established measure of fragility to cascading failure "
        "(Koc et al., 2013). System-reserve redundancy and robustness indices for "
        "structures are well developed (Frangopol and Curley, 1987; Ghosn et al., "
        "2010), from which R_S differs by measuring the entropy of the elastic "
        "strain-energy distribution rather than a reserve- or reliability-capacity "
        "ratio. Strain-energy member-importance coefficients are also standard "
        "(Lin et al., 2019). Against this background the present index is a "
        "specific recombination rather than a new principle: it is the entropy of "
        "the member-level elastic strain-energy distribution, evaluated under "
        "code-style single-member removal, on planar frames. What is genuinely new "
        "is the evidence in the sections below that this particular construction "
        "(i) tracks the degree of static indeterminacy, (ii) is not reducible to "
        "compliance-based importance, and (iii) reaches the collapse onset with no "
        "calibration, placed side by side with the standard criteria on one "
        "analysis.", "Paragraph")
    P(doc, "Two caveats belong here at the start. Strain-energy entropy is itself "
        "old, so what is new is the robustness formulation and the evidence, not "
        "the use of entropy. And I make no claim of early warning. Work on "
        "feedback-amplified systems has shown that entropy collapse can be a "
        "first-order transition with no statistical precursor, so the entropy drop "
        "seen here is read as a marker of localization rather than a forecast of "
        "it.", "Paragraph")

    # ---- 2 Materials and methods ----
    P(doc, "2 MATERIALS AND METHODS", "Heading1")
    P(doc, "2.1 strain-energy entropy", "Heading2")
    P(doc, "At a given load, let U_i be the elastic strain energy in active member "
        "i, with N active members in all. Equation (1) gives the normalized energy "
        "distribution and its Shannon entropy.", "Paragraph")
    equation(doc, "p_i = U_i / sum(U_j),   S = - sum p_i ln(p_i),   H = S / ln(N)", 1)
    P(doc, "H equals 1 when the energy is shared evenly and 0 when a single member "
        "holds it all. Because the fractions p_i are ratios, S and H do not move "
        "as the load is scaled on a fixed topology. They shift only when the load "
        "path itself changes.", "Paragraph")

    P(doc, "2.2 alternate load path and the robustness index", "Heading2")
    P(doc, "For a linear-elastic frame the redistribution after a member is lost "
        "follows exactly from dropping that member from the stiffness assembly and "
        "solving Ku = F again; no separate redistribution rule is needed. If the "
        "reduced stiffness matrix loses positive-definiteness, the removal has "
        "produced a mechanism. Writing H_0 for the intact normalized entropy and "
        "H_k for the survivors' entropy after removing member k (with H_k = 0 for a "
        "mechanism), the entropy drop is dH_k = H_0 - H_k and the index is the "
        "mean over the removal set R, Equation (2).", "Paragraph")
    equation(doc, "R_S = (1/|R|) sum_k H_k,   in [0, 1]", 2)
    P(doc, "An R_S near 1 means the frame stays redundant under any single loss; "
        "near 0 means it does not. The index is bounded and needs no threshold "
        "tuned to the structure. Two summaries accompany it: the worst case "
        "min_k H_k, and R_S restricted to the removals that stay stable, which "
        "separates how even the surviving distribution is from how many losses are "
        "survivable. A determinacy bound makes the link to redundancy precise: a "
        "statically determinate truss has R_S = 0, because removing any bar of it "
        "forms a mechanism, so R_S > 0 requires at least one degree of static "
        "indeterminacy (proof in the accompanying theory note).", "Paragraph")

    P(doc, "2.3 four criteria for comparison", "Heading2")
    P(doc, "The four criteria are read off one incremental analysis. Load is "
        "scaled up, over-stressed members drop out, and at every step four "
        "quantities are recorded; each criterion fires the first time its quantity "
        "crosses. Displacement fires when the largest nodal translation passes a "
        "drift limit. DCR fires when the highest demand-capacity ratio reaches "
        "yield. The energy criterion watches the compliance U/lambda^2 rather than "
        "U itself, since U grows with lambda^2 on a fixed topology and would "
        "otherwise trip on load alone. The entropy criterion fires when the "
        "step-to-step entropy drop is an outlier under a z-score test that looks "
        "only at earlier steps. The first three need a threshold chosen for the "
        "structure. The entropy test does not.", "Paragraph")

    P(doc, "2.4 implementation and verification", "Heading2")
    P(doc, "The tool is written in Python, with modules that talk to each other "
        "only through shared data classes. The frame element is a planar "
        "Euler-Bernoulli element with axial and in-plane bending stiffness; truss "
        "members are axial-only bars. One detail matters more than it looks: the "
        "local bending plane is pinned to the global X-Y plane for every member. "
        "Taking it from a per-member reference vector rotates a column's bending "
        "plane out of plane and leaves the in-plane system singular, a failure "
        "that is easy to miss because a least-squares solve still returns "
        "numbers.", "Paragraph")
    P(doc, "Verification runs on three tiers (Table 1). The first two confirm the "
        "solver: closed-form beams match to the last digit, and an independent "
        "three-degree-of-freedom solver written from scratch reproduces the "
        f"redundant frames exactly, including a larger 3-bay 6-story frame "
        f"({R['large_n'][0]} nodes, {R['large_n'][1]} members). The third tier "
        "checks the quantities R_S actually depends on: the per-member "
        "distribution, the post-removal ALP states, the failure criterion, and the "
        "stability test all agree with independent calculations. The Ziemian and "
        "Ziemian (2021) steel benchmark frames are a useful external reference, but "
        "their published results are second-order (P-Delta) and sit outside what a "
        "first-order solver can reproduce, so I cite them for future checks rather "
        "than match them here.", "Paragraph")
    ver_rows = []
    for r in R["analytical"]:
        ver_rows.append(["Analytical", r["case"], f"{r['delta_err_pct']:.4f}%",
                         f"{r['U_err_pct']:.4f}%"])
    for r in R["independent"]:
        ver_rows.append(["Independent", r["frame"][:34], f"{r['disp_err_pct']:.4f}%",
                         f"{r['U_err_pct']:.4f}%"])
    n_pass = sum(1 for c in R["index_checks"] if c["pass"])
    ver_rows.append(["Index checks",
                     f"distribution / ALP / failure / stability "
                     f"({n_pass}/{len(R['index_checks'])} pass)", "-", "-"])
    table(doc, "Table 1 Verification errors relative to reference (computed by "
               "benchmark.py).",
          ["Tier", "Case", "Displacement err", "Strain-energy err"], ver_rows)

    # ---- 3 Results ----
    P(doc, "3 RESULTS", "Heading1")
    P(doc, "3.1 robustness index", "Heading2")
    P(doc, "Three frames are analysed: a two-span beam with no redundancy, a 2-bay "
        "3-story steel moment frame, and a 6-panel X-braced truss bridge. Table 2 "
        "collects the index for each.", "Paragraph")
    table(doc, "Table 2 Entropy Robustness Index of the benchmark frames (computed "
               "by entropy/robustness.py).",
          ["Frame", "Intact H0", "R_S", "Worst-case H_k", "Mechanism fraction"],
          [["Two-span beam", f3(R["beam"].intact_entropy_norm),
            f3(R["beam"].robustness_index), f3(R["beam"].worst_case_entropy_norm),
            f"{R['beam'].mechanism_fraction:.0%}"],
           ["Moment frame (2-bay, 3-story)", f3(R["frame"].intact_entropy_norm),
            f3(R["frame"].robustness_index), f3(R["frame"].worst_case_entropy_norm),
            f"{R['frame'].mechanism_fraction:.0%}"],
           ["Truss bridge (X-braced)", f3(R["truss"].intact_entropy_norm),
            f3(R["truss"].robustness_index), f3(R["truss"].worst_case_entropy_norm),
            f"{R['truss'].mechanism_fraction:.0%}"]])
    P(doc, "The values fall where intuition says they should. The two-span beam "
        "has no second path: lose either member and what remains is a mechanism, "
        f"so R_S = 0. The moment frame and the X-braced truss both survive every "
        f"single removal (R_S = {rs_frame} and {rs_truss}), but the worst single "
        f"loss in the truss pulls its surviving distribution down to H_k = "
        f"{f3(R['truss'].worst_case_entropy_norm)}, a sharper localization than "
        "anything in the moment frame. The larger 3-bay 6-story frame behaves the "
        f"same way (R_S = {f2(R['large'].robustness_index)}), so the index is not "
        "an artefact of small models.", "Paragraph")
    figure(doc, os.path.join(FIG, "fig_robustness_ranking.png"),
           "Figure 1 Entropy-based member criticality (entropy drop dH_k) for the "
           "moment frame and the truss bridge.")

    P(doc, "3.2 R_S tracks redundancy", "Heading2")
    sweep = R["sweep"]
    P(doc, "To test whether R_S measures redundancy rather than merely re-packaging "
        "stiffness, I sweep the truss bridge from its statically determinate "
        "single-diagonal form to its fully X-braced form, adding one "
        "counter-diagonal at a time. The degree of static indeterminacy is then "
        "known exactly. R_S rises monotonically from "
        f"{f2(sweep[0].robustness_index)} at zero redundancy to "
        f"{f2(sweep[-1].robustness_index)} at six, while the fraction of "
        "single-loss mechanisms falls from 100% to 0% (Figure 2, Table 3). The "
        "determinate case gives exactly R_S = 0, as the determinacy bound "
        "requires.", "Paragraph")
    table(doc, "Table 3 R_S versus degree of static indeterminacy (DSI) for the "
               "truss bridge (computed by analysis/parametric.py).",
          ["DSI", "Members", "R_S", "Mechanism fraction"],
          [[str(p.dsi), str(p.n_members), f3(p.robustness_index),
            f"{p.mechanism_fraction:.0%}"] for p in sweep])
    figure(doc, os.path.join(FIG, "fig_redundancy_sweep.png"),
           "Figure 2 R_S rises and the single-loss mechanism fraction falls as the "
           "truss gains redundancy.")

    P(doc, "3.3 the ranking is not compliance importance", "Heading2")
    P(doc, "A natural objection is that dH_k just restates the compliance increase "
        "caused by removing a member. The two are related but not the same. "
        "Comparing the entropy drop with the removal-based compliance importance "
        f"I_k = dU/U_0 gives a Spearman rank correlation of {rho_f} for the moment "
        f"frame and {rho_t} for the truss (Figure 3). The correlation is "
        "frame-dependent and well below one; for the moment frame the two rankings "
        "are uncorrelated to weakly negative, so the member whose loss most softens "
        "the structure need not be the one whose loss most localizes the energy. I "
        "do not claim statistical significance at this sample size; the point is "
        "that dH_k is not a monotone re-encoding of compliance importance.",
      "Paragraph")
    figure(doc, os.path.join(FIG, "fig_importance_scatter.png"),
           "Figure 3 Entropy criticality dH_k against compliance importance I_k; "
           "the two rankings differ (Spearman in titles).")
    ens = R["ensemble"]
    rhos = [e.spearman_rho for e in ens]
    n_neg = sum(1 for r in rhos if r < 0)
    P(doc, "To check that this is a property of the measure and not of one example, "
        f"I repeat the comparison over an ensemble of {len(ens)} frames of varied "
        "size and type (moment frames from one to four bays and two to six "
        "stories, and redundant trusses). The rank correlation ranges from "
        f"{min(rhos):+.2f} to {max(rhos):+.2f} with mean "
        f"{sum(rhos)/len(rhos):+.2f}, and is negative for {n_neg} of the "
        f"{len(ens)} frames (Figure 4). The sign is systematic rather than random: "
        "every multi-bay moment frame gives a negative correlation, while small "
        "frames and trusses give positive ones. The entropy criticality is "
        "therefore consistently far from a monotone re-encoding of compliance "
        "importance, and for a whole class of structures it ranks members in "
        "essentially the opposite order.", "Paragraph")
    figure(doc, os.path.join(FIG, "fig_importance_ensemble.png"),
           "Figure 4 Spearman correlation between entropy criticality and "
           "compliance importance across the frame ensemble; it stays well below 1 "
           "and is negative for the multi-bay moment frames.")

    P(doc, "3.4 progressive collapse and criteria comparison", "Heading2")
    tr = R["triggers"]
    P(doc, "Loading the truss bridge step by step to failure shows the entropy "
        "behaving as expected. While the truss is intact the entropy holds flat, "
        "since it ignores load magnitude; at the first member failure it drops as "
        "the energy crowds into fewer members, then recovers part of the way as "
        "the remaining bars take up the load (Figure 5). Table 4 lists the load "
        "factor at which each criterion first fires.", "Paragraph")
    figure(doc, os.path.join(FIG, "fig_entropy_trajectory.png"),
           "Figure 5 Normalized entropy and compliance versus load step for the "
           "truss bridge; both are flat until the first member failure, then "
           "respond together.")
    table(doc, "Table 4 First-trigger load factor of the four criteria (truss "
               "bridge, load-step 0.3; computed by analysis/criteria.py).",
          ["Criterion", "First trigger (load factor)", "Threshold required"],
          [["Displacement", f2(tr["displacement"]), "yes (drift limit)"],
           ["DCR", f2(tr["dcr"]), "yes (acceptance DCR)"],
           ["Energy (compliance)", f2(tr["energy"]), "yes (energy multiple)"],
           ["Entropy", f2(tr["entropy"]), "no"]])
    figure(doc, os.path.join(FIG, "fig_criteria_comparison.png"),
           "Figure 6 First-trigger load factor for the four criteria; grey bars "
           "need a calibrated threshold, the entropy bar does not.")
    P(doc, "The displacement criterion trips first and where it trips is set by the "
        "drift limit, a known soft spot of displacement-based checks. The other "
        f"three agree near load factor {f2(tr['dcr'])} to {f2(tr['energy'])}. DCR "
        "catches the stress that triggers the first failure; the energy and "
        "entropy criteria catch its aftermath at the same load. I want to be "
        "precise about this: in the linear-elastic range the entropy is flat until "
        "the first removal, so the entropy criterion is a one-step-lagging, "
        "consequence-side indicator, not an onset detector that leads DCR. A "
        "step-size study confirms it: the DCR-to-entropy gap is one load step at "
        "every step size and shrinks with it. The honest claim is that entropy "
        "reaches the same conclusion as DCR and energy without a calibrated "
        "threshold, not that it sees collapse sooner.", "Paragraph")

    P(doc, "3.5 limitations", "Heading2")
    P(doc, "The limits should be read plainly. Entropy of strain energy is old, so "
        "the new part is the robustness formulation and the evidence that it tracks "
        "redundancy, departs from compliance importance, and is calibration-free. "
        "The analysis is planar, linear-elastic, first-order, and quasi-static, "
        "which leaves out plasticity, second-order and catenary action, "
        "out-of-plane response, dynamic amplification, and connection failure. "
        "Linear-static ALP is an accepted and usually conservative code method, "
        "though for stability-governed compression members the first-order "
        "assumption can be unconservative. And the entropy drop lands with the "
        "localization rather than ahead of it, so the claim is a calibration-free "
        "measure of robustness and localization, not lead time.", "Paragraph")

    # ---- 4 Conclusion ----
    P(doc, "4 CONCLUSION", "Heading1")
    P(doc, "R_S is a bounded redundancy measure, built from the strain-energy "
        "entropy under the alternate-load-path procedure, that carries no "
        "calibrated threshold and equals 0 for a statically determinate structure "
        "by a provable bound. On a solver verified to floating-point precision "
        "(closed-form, an independent solver, and checks of the distribution, ALP "
        "states, failure criterion and stability test), R_S separates the "
        f"benchmarks ({rs_beam}, {rs_frame}, {rs_truss}) and rises monotonically "
        "with the degree of static indeterminacy. The entropy member-criticality "
        f"ranking is not a monotone re-encoding of compliance importance (Spearman "
        f"{rho_f} and {rho_t}), and the entropy criterion reaches the collapse "
        "onset at the same load as the resistance and energy criteria with no "
        "threshold tuned to the structure.", "Paragraph")

    # ---- Back matter ----
    P(doc, "Author's Contributions: LAJSS endorses the CRediT taxonomy of "
           "contributor roles. Conceptualization, F Carvajal Brown; Methodology, "
           "F Carvajal Brown; Software, F Carvajal Brown; Validation, F Carvajal "
           "Brown; Formal analysis, F Carvajal Brown; Investigation, F Carvajal "
           "Brown; Writing - original draft, F Carvajal Brown; Writing - review & "
           "editing, F Carvajal Brown.", "fnAuthor-notesContributions")
    P(doc, "Data Availability: Research data is available in a repository. The "
           "source code, scenarios, and scripts that reproduce every number and "
           "figure in this article are openly available at "
           "https://github.com/fcarvajalbrown/Entropy-Collapse-Simulator "
           "(GPL-3.0-or-later).", "dataStatement")
    P(doc, "Conflicts of Interest: The author declares no conflict of interest.",
      "dataStatement")
    P(doc, "Funding: This research received no external funding.", "dataStatement")
    P(doc, "Editor: (Filled by: Journal Staff)", "fnEdited-by")

    # ---- References ----
    P(doc, "References", "ReferenceHead")
    for ref in REFERENCES:
        P(doc, ref, "Reference")

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


# LAJSS reference style: Surname, Initials. (year). Title. Journal Volume:pages.
REFERENCES = [
    "Department of Defense (2016). Unified Facilities Criteria UFC 4-023-03: "
    "Design of Buildings to Resist Progressive Collapse, US Department of Defense.",
    "Feng, D.C., et al. (2024). Physically-based collapse failure criteria in "
    "progressive collapse analyses of random-parameter multi-story RC structures "
    "subjected to column removal scenarios. Engineering Structures. "
    "https://doi.org/10.1016/j.engstruct.2024.119412",
    "Frangopol, D.M., Curley, J.P. (1987). Effects of damage and redundancy on "
    "structural reliability. Journal of Structural Engineering 113(7):1533-1549.",
    "General Services Administration (2003). Progressive Collapse Analysis and "
    "Design Guidelines for New Federal Office Buildings and Major Modernization "
    "Projects, US General Services Administration.",
    "Ghosn, M., Moses, F., Frangopol, D.M. (2010). Redundancy and robustness of "
    "highway bridge superstructures and substructures. Structure and "
    "Infrastructure Engineering 6(1-2):257-278.",
    "Koc, Y., Warnier, M., Kooij, R.E., Brazier, F.M. (2013). An entropy-based "
    "metric to quantify the robustness of power grids against cascading failures. "
    "Safety Science 59:126-134.",
    "Lin, K., et al. (2019). Importance assessment of structural members based on "
    "elastic-plastic strain energy. Advances in Materials Science and Engineering "
    "2019:8019675.",
    "Nafday, A.M. (2011). Consequence-based structural design approach for black "
    "swan events. Structural Safety 33(1):108-114.",
    "Shannon, C.E. (1948). A mathematical theory of communication. Bell System "
    "Technical Journal 27:379-423.",
    "von Scheven, M., Ramm, E., Bischoff, M. (2021). Quantification of the "
    "redundancy distribution in truss and beam structures. International Journal "
    "of Solids and Structures 213:41-49. "
    "https://doi.org/10.1016/j.ijsolstr.2020.11.002",
    "Ziemian, C.W., Ziemian, R.D. (2021). Steel benchmark frames for structural "
    "analysis and validation studies: Finite element models and numerical "
    "simulation data. Data in Brief 39:107564. "
    "https://doi.org/10.1016/j.dib.2021.107564",
    "Ziha, K. (2000). Redundancy and robustness of systems of events. "
    "Probabilistic Engineering Mechanics 15(4):347-357.",
]


if __name__ == "__main__":
    build()
